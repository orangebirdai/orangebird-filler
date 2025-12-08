import os
import uuid
import io
import json
import re
from fastapi import FastAPI, File, UploadFile, Form, Request
from fastapi.responses import FileResponse, HTMLResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from docx import Document
from docx.shared import Pt, RGBColor
from PyPDF2 import PdfReader
from groq import Groq
from dotenv import load_dotenv
# ←←← ADD THESE TWO LINES AT THE VERY TOP OF app.py (with your other imports)
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
load_dotenv()
client = Groq(api_key=os.getenv("GROQ_API_KEY"))

app = FastAPI()
app.mount("/static", StaticFiles(directory="static"), name="static")
templates = Jinja2Templates(directory="templates")
os.makedirs("uploads", exist_ok=True)

def extract_text(content: bytes, name: str) -> str:
    if name.lower().endswith(".pdf"):
        return "\n".join(p.extract_text() or "" for p in PdfReader(io.BytesIO(content)).pages)
    doc = Document(io.BytesIO(content))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

def count_words(text: str) -> int:
    return len(re.findall(r'\b\w+\b', text))

# ←←← REPLACE YOUR ENTIRE make_docx FUNCTION WITH THIS ONE
def make_docx(md: str, path: str):
    doc = Document()
    title_added = False
    current_heading = None

    for raw_line in md.split("\n"):
        line = raw_line.rstrip()
        if not line:
            doc.add_paragraph("")
            continue

        if not title_added and line.startswith("# "):
            p = doc.add_paragraph()
            p.add_run(line[2:]).bold = True
            p.style = "Title"
            title_added = True
            continue

        if line.startswith("## "):
            current_heading = line[3:].strip()
            p = doc.add_paragraph()
            p.add_run(current_heading).bold = True
            p.paragraph_format.space_after = Pt(6)
            continue

        if current_heading and line.strip() == current_heading:
            continue

        p = doc.add_paragraph()

        # Find and make real clickable hyperlinks
        url_pattern = re.compile(r'(https?://[^\s]+|doi\.org/[^\s]+)')
        last_end = 0
        for match in url_pattern.finditer(line):
            start, end = match.span()
            if start > last_end:
                p.add_run(line[last_end:start])

            url = match.group(0)
            if url.startswith("doi.org"):
                url = "https://" + url

            run = p.add_run(url)
            run.font.color.rgb = RGBColor(0, 0, 255)
            run.underline = True

            # This makes it ACTUALLY clickable in Word
            r_id = doc.part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('r:id'), r_id)
            new_run = OxmlElement('w:r')
            new_run.append(run._element)
            hyperlink.append(new_run)
            p._p.append(hyperlink)

            last_end = end

        if last_end < len(line):
            p.add_run(line[last_end:])

    doc.save(path)

@app.get("/", response_class=HTMLResponse)
async def home(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})

@app.post("/go")
async def go(file: UploadFile = File(...), style: str = Form("MLA"), hint: str = Form(""), wordcount: str = Form("1500")):
    raw_text = extract_text(await file.read(), file.filename)
    try:
        target_words = max(800, min(7000, int(wordcount)))
    except:
        target_words = 1500

    # 1. Worksheet
    prompt1 = f"""Complete this worksheet using {style}. Answer every question in order.
Topic hint: {hint or 'none'}.

WORKSHEET:
\"\"\"{raw_text}\"\"\"

Return ONLY clean markdown."""
    resp1 = client.chat.completions.create(model="llama-3.3-70b-versatile", messages=[{"role": "user", "content": prompt1}], temperature=0.2, max_tokens=8000)
    worksheet_md = resp1.choices[0].message.content
    worksheet_file = f"COMP_{uuid.uuid4().hex[:10]}.docx"
    make_docx(worksheet_md, f"uploads/{worksheet_file}")

    # 2. Real sources
    sources_prompt = f"""Give exactly 8 real peer-reviewed articles about this commodity.
JSON only: [{{"author":"Last, First","title":"...","journal":"...","year":"2024","doi":"https://doi.org/..."}}]"""
    sources_resp = client.chat.completions.create(model="llama-3.3-70b-versatile", messages=[{"role": "user", "content": sources_prompt + f"\n\nWORKSHEET:\n\"\"\"{worksheet_md}\"\"\""}], temperature=0.3, max_tokens=4000)
    try:
        sources = json.loads(sources_resp.choices[0].message.content)
    except:
        sources = []

    bib_heading = "Works Cited" if style.upper() != "APA" and style.upper() != "CHICAGO" else ("References" if style.upper() == "APA" else "Bibliography")
    bib_lines = [f"{bib_heading}\n"]
    for s in sources[:8]:
        entry = f"{s.get('author','Unknown')}. \"{s.get('title','Untitled')}.\" {s.get('journal','')}, {s.get('year','n.d.')}"
        if s.get("doi"): entry += f", {s.get('doi')}"
        entry += "."
        bib_lines.append(entry)
        bib_lines.append("")
    works_cited = "\n".join(bib_lines)

    # 3. Title + outline
    outline_prompt = "Create a strong title and 8–12 section outline. Return ONLY JSON: {\"title\": \"...\", \"outline\": [...]}\""
    outline_resp = client.chat.completions.create(model="llama-3.3-70b-versatile", messages=[{"role": "user", "content": outline_prompt + f"\n\nWORKSHEET:\n\"\"\"{worksheet_md}\"\"\""}], temperature=0.3, max_tokens=2000)
    try:
        plan = json.loads(outline_resp.choices[0].message.content)
    except:
        plan = {"title": "The Real Story of This Commodity", "outline": [f"Section {i}" for i in range(1,11)]}

    # 4. REAL essay — no placeholder
    full_essay = f"# {plan['title']}\n\n"
    current_words = 0

    for i, heading in enumerate(plan["outline"], 1):
        if current_words >= target_words:
            break
        remaining = target_words - current_words
        words_this_section = min(750, remaining + 150)

        section_prompt = f"""Write ONLY body text for section titled exactly: {heading}
Target ~{words_this_section} words.
55-year-old American senior analyst voice — first-person or “we/you”, contractions, casual markers, bursty sentences.
Use proper {style} in-text citations. No repetition.

WORKSHEET:
\"\"\"{worksheet_md}\"\"\"

{bib_heading}:
{works_cited}

Return ONLY clean markdown body."""
        resp = client.chat.completions.create(model="llama-3.3-70b-versatile", messages=[{"role": "user", "content": section_prompt}], temperature=0.5, max_tokens=3000)
        section_text = resp.choices[0].message.content.strip()
        section_words = count_words(section_text)
        full_essay += f"## {heading}\n\n{section_text}\n\n"
        current_words += section_words

    full_essay += works_cited
    essay_file = f"ESSAY_{uuid.uuid4().hex[:10]}.docx"
    make_docx(full_essay, f"uploads/{essay_file}")

    # SUCCESS PAGE
        # FINAL SUCCESS PAGE — PERFECT SPACING
    return HTMLResponse(f"""
    <!DOCTYPE html>
    <html>
    <head>
        <title>OrangeBird — DONE!</title>
        <style>
            body {{ font-family: Arial; text-align: center; padding: 100px 20px; background: #2c3e50; color: white; }}
            h1 {{ font-size: 70px; color: #f1c40f; margin-bottom: 40px; }}
            .btn {{ 
                display: block; width: 380px; max-width: 90%; margin: 30px auto; 
                padding: 22px; font-size: 26px; font-weight: bold; 
                border-radius: 15px; text-decoration: none; color: white;
            }}
            .btn-worksheet {{ background: #e67e22; }}
            .btn-worksheet:hover {{ background: #d35400; }}
            .btn-essay {{ background: #2980b9; }}
            .btn-essay:hover {{ background: #1a6ea3; }}
            .back {{ margin-top: 80px; font-size: 20px; color: #bdc3c7; text-decoration: none; }}
        </style>
    </head>
    <body>
        <h1>DONE!</h1>
        <p style="font-size:24px;">Your files are ready!</p>
        <a href="/download/{worksheet_file}" download class="btn btn-worksheet">Download Completed Worksheet</a>
        <a href="/download/{essay_file}" download class="btn btn-essay">Download Essay ({target_words} words)</a>
        <br><br><br>
        <a href="/" class="back">← Generate Another</a>
    </body>
    </html>
    """)

@app.get("/download/{filename:path}")
async def download(filename: str):
    file_path = f"uploads/{filename}"
    if not os.path.exists(file_path):
        return HTMLResponse("File not found — please generate again.", status_code=404)
    return FileResponse(file_path, filename=filename)
